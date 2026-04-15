# tests/test_word_output.py
import pytest
from pathlib import Path
from unittest.mock import patch


@pytest.fixture(autouse=True)
def isolated_projects_dir(tmp_path, monkeypatch):
    monkeypatch.setenv("PROJECTS_DIR", str(tmp_path))
    from api.config import get_settings
    get_settings.cache_clear()
    yield tmp_path
    get_settings.cache_clear()


@pytest.fixture
def slug(isolated_projects_dir):
    slug = "word-output-test"
    outputs_dir = isolated_projects_dir / slug / "outputs"
    outputs_dir.mkdir(parents=True)
    return slug


@pytest.fixture
def sample_sections():
    return [
        {"title": "Executive Summary", "content": "This is the executive summary."},
        {"title": "Case for Change", "content": "The case for change is compelling."},
    ]


@pytest.fixture
def sample_metadata():
    return {
        "org_name": "Acme Logistics Ltd",
        "financial_year": "FY2026",
        "sponsor": "Jane Smith, CEO",
    }


def test_word_output_writes_file(slug, sample_sections, sample_metadata):
    """WordOutputTool creates a .docx file at the returned path."""
    from agents.tools.word_output import WordOutputTool
    with patch("agents.tools.word_output.insert_agent_output_sync"):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="business_plan.docx",
            agent_name="business_plan_generator",
        )
    assert Path(result).exists()


def test_word_output_returns_absolute_path(slug, sample_sections, sample_metadata):
    """Return value is an absolute path ending in .docx."""
    from agents.tools.word_output import WordOutputTool
    with patch("agents.tools.word_output.insert_agent_output_sync"):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="business_plan.docx",
            agent_name="business_plan_generator",
        )
    assert Path(result).is_absolute()
    assert result.endswith(".docx")


def test_word_output_contains_section_titles(slug, sample_sections, sample_metadata):
    """Generated .docx contains each section title as a heading."""
    from agents.tools.word_output import WordOutputTool
    import docx as python_docx
    with patch("agents.tools.word_output.insert_agent_output_sync"):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="business_plan.docx",
            agent_name="business_plan_generator",
        )
    doc = python_docx.Document(result)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in full_text
    assert "Case for Change" in full_text


def test_word_output_contains_org_name(slug, sample_sections, sample_metadata):
    """org_name from metadata appears in the document."""
    from agents.tools.word_output import WordOutputTool
    import docx as python_docx
    with patch("agents.tools.word_output.insert_agent_output_sync"):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="business_plan.docx",
            agent_name="business_plan_generator",
        )
    doc = python_docx.Document(result)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Acme Logistics Ltd" in full_text


def test_word_output_appends_docx_extension(slug, sample_sections, sample_metadata):
    """filename without .docx extension gets it added automatically."""
    from agents.tools.word_output import WordOutputTool
    with patch("agents.tools.word_output.insert_agent_output_sync"):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="business_plan_no_ext",
            agent_name="business_plan_generator",
        )
    assert result.endswith(".docx")
    assert Path(result).exists()


def test_word_output_error_on_write_failure(slug, sample_sections, sample_metadata):
    """Returns error string when the file cannot be saved."""
    from agents.tools.word_output import WordOutputTool
    with patch("agents.tools.word_output.insert_agent_output_sync"), \
         patch("docx.document.Document.save", side_effect=OSError("disk full")):
        tool = WordOutputTool(slug=slug)
        result = tool._run(
            sections=sample_sections,
            metadata=sample_metadata,
            filename="fail.docx",
            agent_name="business_plan_generator",
        )
    assert result.startswith("Error: render failed")
