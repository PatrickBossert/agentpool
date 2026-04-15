# agents/tools/word_output.py
from datetime import date
from pathlib import Path
from typing import Any
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
from api.config import get_settings
from agents.tools._db import insert_agent_output_sync


class WordOutputToolInput(BaseModel):
    sections: list[dict[str, str]] = Field(
        description=(
            "List of dicts with 'title' and 'content' keys — one per business plan section."
        )
    )
    metadata: dict[str, Any] = Field(
        description=(
            "Dict with keys: org_name, financial_year, sponsor. "
            "Optional: prepared_by, date (auto-filled if absent)."
        )
    )
    filename: str = Field(
        description="Output filename (e.g. 'business_plan.docx'). "
        ".docx extension added automatically if missing."
    )
    agent_name: str = Field(
        description="Name of the agent producing this output (used for output tracking)."
    )


class WordOutputTool(BaseTool):
    name: str = "WordOutputTool"
    description: str = (
        "Write a business plan to a styled .docx file in the project outputs directory. "
        "Pass sections as a list of {title, content} dicts, metadata with org_name/"
        "financial_year/sponsor, and a filename. Returns the absolute file path."
    )
    args_schema: type[BaseModel] = WordOutputToolInput
    slug: str

    def _run(
        self,
        sections: list[dict[str, str]],
        metadata: dict[str, Any],
        filename: str,
        agent_name: str,
    ) -> str:
        try:
            import docx
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return "Error: python-docx not installed — run: pip install python-docx"

        settings = get_settings()
        outputs_dir = Path(settings.projects_dir) / self.slug / "outputs"
        outputs_dir.mkdir(parents=True, exist_ok=True)

        if not filename.endswith(".docx"):
            filename = f"{filename}.docx"
        file_path = outputs_dir / filename

        try:
            doc = docx.Document()

            # ── Title page ────────────────────────────────────────────────────
            _navy = RGBColor(0x1F, 0x39, 0x7D)

            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(metadata.get("org_name", ""))
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = _navy

            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = subtitle_para.add_run("Digital Modernisation Business Plan")
            run2.font.size = Pt(18)
            run2.font.color.rgb = _navy

            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fy = metadata.get("financial_year", "")
            sponsor = metadata.get("sponsor", "")
            doc_date = metadata.get("date", str(date.today()))
            meta_para.add_run(f"{fy}  |  {sponsor}  |  {doc_date}")

            doc.add_page_break()

            # ── Table of contents ─────────────────────────────────────────────
            toc_heading = doc.add_heading("Contents", level=1)
            for h_run in toc_heading.runs:
                h_run.font.color.rgb = _navy
            for i, section in enumerate(sections, start=1):
                doc.add_paragraph(f"{i}.  {section['title']}")

            doc.add_page_break()

            # ── Sections ──────────────────────────────────────────────────────
            for section in sections:
                heading = doc.add_heading(section["title"], level=1)
                for h_run in heading.runs:
                    h_run.font.color.rgb = _navy
                doc.add_paragraph(section.get("content", ""))

            doc.save(file_path)
            insert_agent_output_sync(
                slug=self.slug,
                agent_name=agent_name,
                output_type="docx",
                file_path=str(file_path),
            )
        except (OSError, ValueError) as e:
            return f"Error: render failed — {e}"

        return str(file_path)
